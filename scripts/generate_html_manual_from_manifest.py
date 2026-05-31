from __future__ import annotations

import argparse
import base64
import html
import io
import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


MAX_SNIPPET_LINES = 32


INCLUDE_EXTENSIONS = {
    ".java", ".kt", ".py", ".js", ".jsx", ".ts", ".tsx", ".css", ".scss", ".less",
    ".html", ".vue", ".xml", ".properties", ".sql", ".json", ".yaml", ".yml",
}
EXCLUDED_DIRS = {
    ".git", ".idea", ".vscode", "node_modules", "dist", "build", "target", "coverage",
    "__pycache__", ".next", ".nuxt", "软件著作权申请资料",
}


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def project_line_count(project_root: Path) -> int:
    total = 0
    for path in project_root.rglob("*"):
        if path.is_file() and path.suffix.lower() in INCLUDE_EXTENSIONS:
            if not any(part in EXCLUDED_DIRS for part in path.relative_to(project_root).parts):
                total += len(read_text(path).splitlines())
    return total


def data_uri(path: Path) -> str:
    mime = "image/png"
    return f"data:{mime};base64,{base64.b64encode(path.read_bytes()).decode('ascii')}"


def build_manual_image_bytes(path: Path, *, min_width: int = 2400) -> bytes:
    raw = path.read_bytes()
    try:
        from PIL import Image, ImageFilter, ImageOps
    except Exception:
        return raw

    with Image.open(io.BytesIO(raw)) as image:
        prepared = ImageOps.exif_transpose(image).convert("RGB")
        if prepared.width < min_width and prepared.width > 0:
            scale = min_width / prepared.width
            prepared = prepared.resize(
                (round(prepared.width * scale), round(prepared.height * scale)),
                Image.Resampling.LANCZOS,
            )
        prepared = prepared.filter(ImageFilter.UnsharpMask(radius=1.8, percent=155, threshold=2))
        buffer = io.BytesIO()
        prepared.save(buffer, format="PNG", optimize=True)
        return buffer.getvalue()


def data_uri_from_bytes(payload: bytes) -> str:
    return f"data:image/png;base64,{base64.b64encode(payload).decode('ascii')}"


def code_slice(project_root: Path, ref: dict) -> tuple[str, str]:
    file_path = project_root / ref["file"]
    lines = read_text(file_path).splitlines()
    start = max(1, int(ref["start_line"]))
    requested_end = min(len(lines), int(ref["end_line"]))
    end = min(requested_end, start + MAX_SNIPPET_LINES - 1)
    snippet = "\n".join(lines[start - 1:end])
    return f"{ref['file']}（约第{start}-{end}行）", snippet


def esc(text: str) -> str:
    return html.escape(text or "")


def normalize_sentence(text: str) -> str:
    value = (text or "").strip()
    value = re.sub(r"[。！？；;、，,.\s]+$", "", value)
    return value


def join_sentences(items: list[str], separator: str = "；") -> str:
    cleaned = [normalize_sentence(item) for item in items if normalize_sentence(item)]
    return separator.join(cleaned)


def join_terms(items: list[str]) -> str:
    cleaned = [normalize_sentence(item) for item in items if normalize_sentence(item)]
    return "、".join(cleaned)


def join_field_meanings(fields: list[dict]) -> str:
    pairs = []
    for field in fields:
        name = normalize_sentence(field.get("name", ""))
        meaning = normalize_sentence(field.get("meaning", ""))
        if name or meaning:
            pairs.append(f"{name}:{meaning}".strip(":"))
    return "；".join(pairs)


def infer_feature_categories(manifest: dict) -> str:
    categories = [normalize_sentence(item) for item in manifest.get("technical_feature_categories", []) if normalize_sentence(item)]
    if categories:
        return "、".join(categories[:2])
    system_name = str(manifest.get("system_name", ""))
    industry = str(manifest.get("industry", ""))
    hints = f"{system_name} {industry}"
    if any(token in hints for token in ["医院", "医疗", "诊疗", "病历", "处方"]):
        return "医疗软件、人工智能软件"
    if any(token in hints for token in ["仪器", "控制", "实验室", "设备"]):
        return "物联网软件、人工智能软件"
    if any(token in hints for token in ["教育", "校园", "课堂"]):
        return "教育软件、人工智能软件"
    return "人工智能软件、应用软件"


def build_feature_description(manifest: dict) -> str:
    features = [normalize_sentence(item) for item in manifest.get("technical_features", []) if normalize_sentence(item)]
    if not features:
        return "采用模块化架构、规则引擎和结构化导出机制，支持业务联动、结果展示与资料沉淀。"
    return "；".join(features)


def build_feature_table_text(manifest: dict) -> str:
    categories = infer_feature_categories(manifest)
    features = [normalize_sentence(item) for item in manifest.get("technical_features", []) if normalize_sentence(item)]
    if not features:
        return f"{categories}；支持业务联动、结果展示和结构化导出。"
    head = features[0]
    if len(features) > 1:
        return f"{categories}；{head}。"
    return f"{categories}；{head}。"


def collect_module_screenshots(module: dict) -> list[dict]:
    screenshots: list[dict] = []
    for item in module.get("screenshots", []) or []:
        if isinstance(item, dict) and item.get("file"):
            screenshots.append(item)
    single = module.get("screenshot") or {}
    if isinstance(single, dict) and single.get("file"):
        screenshots.append(single)
    deduped: list[dict] = []
    seen: set[str] = set()
    for shot in screenshots:
        key = str(shot.get("file", "")).strip()
        if key and key not in seen:
            seen.add(key)
            deduped.append(shot)
    return deduped


def detect_project_shape(project_root: Path) -> dict:
    return {
        "has_package_json": (project_root / "package.json").exists(),
        "has_java_server": (project_root / "java-server").exists(),
        "has_run_java_script": (project_root / "scripts" / "run-java.ps1").exists(),
        "has_build_java_script": (project_root / "scripts" / "build-java.ps1").exists(),
    }


def infer_entry_urls(project_root: Path) -> list[str]:
    readme = project_root / "README.md"
    if not readme.exists():
        return []
    text = read_text(readme)
    urls = []
    for line in text.splitlines():
        line = line.strip()
        if line.startswith("- `http://") or line.startswith("- `https://"):
            start = line.find("`") + 1
            end = line.find("`", start)
            if start > 0 and end > start:
                urls.append(line[start:end])
    return urls


def build_install_steps(project_root: Path, manifest: dict) -> tuple[list[str], str]:
    shape = detect_project_shape(project_root)
    urls = infer_entry_urls(project_root)
    architecture = str(manifest.get("architecture", ""))
    is_desktop = "桌面" in architecture or ((project_root / "src" / "main" / "java").exists() and not shape["has_package_json"])
    steps = ["获取项目源代码后进入项目根目录。"]

    if shape["has_package_json"]:
        steps.append("执行 npm install 安装前端依赖。")

    if is_desktop and shape["has_run_java_script"]:
        steps.append(r"执行 powershell -ExecutionPolicy Bypass -File .\scripts\run-java.ps1 启动桌面程序。")
    elif is_desktop and shape["has_build_java_script"]:
        steps.append(r"先执行 powershell -ExecutionPolicy Bypass -File .\scripts\build-java.ps1 编译源码，再运行桌面程序主入口。")
    elif shape["has_run_java_script"]:
        steps.append(r"执行 powershell -ExecutionPolicy Bypass -File .\scripts\run-java.ps1 启动 Java 服务。")
    elif shape["has_java_server"]:
        steps.append(r"先执行 powershell -ExecutionPolicy Bypass -File .\scripts\build-java.ps1 编译源码，再按项目启动命令运行 Java 服务。")

    if is_desktop:
        steps.append("程序启动后直接进入桌面主窗口，在左侧菜单与顶部操作区完成业务处理。")
    elif urls:
        steps.append(f"启动完成后在浏览器中访问 {join_sentences(urls, ' 或 ')} 进入系统页面。")
    else:
        steps.append("启动完成后在浏览器中访问项目配置的本地地址进入系统页面。")

    if is_desktop and shape["has_build_java_script"]:
        warning = r"如需单独校验桌面程序编译状态，可执行 powershell -ExecutionPolicy Bypass -File .\scripts\build-java.ps1 完成 Java 源码编译。"
    elif shape["has_build_java_script"]:
        warning = r"如需单独校验后端编译状态，可执行 powershell -ExecutionPolicy Bypass -File .\scripts\build-java.ps1 完成 Java 源码编译。"
    else:
        warning = ""

    return steps, warning


def build_export_intro(manifest: dict) -> str:
    export_names = [item.get("name", "") for item in manifest.get("exports", []) if item.get("name")]
    if export_names:
        return f"系统通过服务端导出接口输出{join_sentences(export_names, '、')}等结构化文件，可直接用于归档、复核和统计分析。"
    return "系统通过服务端导出接口输出结构化文件，可直接用于归档、复核和统计分析。"


def html_table(headers: list[str], rows: list[list[str]]) -> str:
    head = "".join(f"<th>{esc(h)}</th>" for h in headers)
    body = "".join("<tr>" + "".join(f"<td>{esc(c)}</td>" for c in row) + "</tr>" for row in rows)
    return f"<table><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>"


def add_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:color"), "C8D3E6")
        borders.append(border)
    tc_pr.append(borders)


def set_font(run, size: float, bold: bool = False, font_name: str = "SimSun") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = False


def set_body_paragraph(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0.74)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5


def build_html(project_root: Path, manifest: dict, output_path: Path) -> None:
    system_name = manifest["system_name"]
    version = manifest.get("version", "V1.0")
    env = manifest["environments"]
    output_path.parent.mkdir(parents=True, exist_ok=True)
    install_steps, install_note = build_install_steps(project_root, manifest)
    export_intro = build_export_intro(manifest)

    feature_categories = infer_feature_categories(manifest)
    feature_description = build_feature_description(manifest)
    feature_table_text = build_feature_table_text(manifest)
    env_rows = [
        ["开发该软件的硬件环境", env.get("dev_hardware", "")],
        ["运行的硬件环境", env.get("run_hardware", "")],
        ["开发该软件的操作系统", env.get("dev_os", "")],
        ["软件开发环境/开发工具", env.get("dev_tools", "")],
        ["该软件的运行平台/操作系统", env.get("run_platform", "")],
        ["软件运行支撑环境/支持软件", env.get("support_env", "")],
        ["编程语言", join_terms(env.get("languages", []))],
        ["开发目的", manifest.get("purpose", "")],
        ["面向领域/行业", manifest.get("industry", "")],
        ["软件的主要功能", join_sentences(manifest.get("main_functions", []))],
        ["软件的技术特点", feature_table_text],
    ]

    operation_modules = []
    detail_modules = []
    modules = manifest.get("modules", [])
    for index, module in enumerate(modules, start=1):
        screenshot_html = ""
        shot_blocks = []
        for shot_index, shot in enumerate(collect_module_screenshots(module), start=1):
            shot_path = project_root / shot["file"]
            if shot_path.exists():
                image_bytes = build_manual_image_bytes(shot_path)
                title = shot.get("title", module["title"])
                label = f"图4-{index}-{shot_index}" if shot_index > 1 else f"图4-{index}"
                shot_blocks.append(
                    f"<div class='figure figure--module'><img src='{data_uri_from_bytes(image_bytes)}' alt='{esc(title)}'>"
                    f"<div class='figcaption'>{label} {esc(title)}</div>"
                    f"</div>"
                )
        screenshot_html = "".join(shot_blocks)
        steps_html = "".join(f"<li>{esc(step if isinstance(step, str) else step.get('text', ''))}</li>" for step in module.get("steps", []))
        operation_modules.append(
            f"<section class='module'><h3>4.{index} {esc(module['title'])}</h3>"
            f"<div class='meta'>菜单：{esc(module.get('menu', ''))}　页面路径：{esc(module.get('route', ''))}</div>"
            f"<ol>{steps_html}</ol>{screenshot_html}</section>"
        )

        refs_html = []
        for ref_index, ref in enumerate(module.get("code_refs", []), start=1):
            ref_label, snippet = code_slice(project_root, ref)
            variables = ref.get("variables", [])
            line_explanations = ref.get("line_explanations", [])
            vars_html = ""
            if variables:
                vars_html = html_table(
                    ["变量名", "含义", "取值范围"],
                    [[item.get("name", ""), item.get("meaning", ""), item.get("range", "")] for item in variables],
                )
            line_html = ""
            if line_explanations:
                line_html = "<ul>" + "".join(
                    f"<li><strong>{esc(item.get('lines', ''))}</strong>：{esc(item.get('explanation', ''))}</li>"
                    for item in line_explanations
                ) + "</ul>"
            tip_html = f"<div class='tip'><strong>逐行解释：</strong>{line_html}</div>" if line_html else ""
            refs_html.append(
                f"<div class='code-card'><h4>5.{index}.{ref_index} {esc(ref.get('title', ref.get('function_name', '核心代码')))}</h4>"
                f"<div class='meta'>{esc(ref_label)}　函数：{esc(ref.get('function_name', ''))}</div>"
                f"<div class='note'>{esc(ref.get('explanation', ''))}</div>"
                f"<pre><code>{esc(snippet)}</code></pre>"
                f"<p><strong>输入：</strong>{esc(ref.get('input', ''))}</p>"
                f"<p><strong>输出：</strong>{esc(ref.get('output', ''))}</p>"
                f"{vars_html}"
                f"{tip_html}"
                f"</div>"
            )

        detail_modules.append(
            f"<section class='module'><h3>5.{index} {esc(module['title'])}</h3>"
            f"<p><strong>功能概述：</strong>{esc(module.get('summary', ''))}</p>"
            f"<p><strong>算法原理：</strong>{esc(module.get('algorithm', ''))}</p>"
            f"<p><strong>输入输出说明：</strong>{esc(module.get('input_output', ''))}</p>"
            f"<p><strong>数据流：</strong>{esc(module.get('data_flow', ''))}</p>"
            f"<p><strong>特殊处理：</strong>{esc(module.get('special_handling', ''))}</p>"
            f"{''.join(refs_html)}</section>"
        )

    export_rows = []
    for item in manifest.get("exports", []):
        export_rows.append([item.get("name", ""), join_field_meanings(item.get("fields", [])), item.get("usage", "")])
    faq_rows = [[item.get("problem", ""), item.get("reason", ""), item.get("steps", ""), item.get("solution", "")] for item in manifest.get("faq", [])]

    test_cases = manifest.get("tests", {}).get("cases", [])
    test_case_rows = [[f"TC-{index:02d}", case] for index, case in enumerate(test_cases, start=1)]
    maintenance_items = manifest.get("maintenance", {}).get("items", [])
    maintenance_html = "".join(
        f"<div class='note'><p><strong>{esc(item.get('title', '维护事项'))}</strong></p>"
        f"<p><strong>对应位置：</strong>{esc(item.get('target', ''))}</p>"
        f"<p><strong>处理方法：</strong>{esc(item.get('method', ''))}</p>"
        f"<p><strong>完成后验证：</strong>{esc(item.get('verify', ''))}</p></div>"
        for item in maintenance_items
    )

    html_doc = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{esc(system_name)} 详细使用说明书</title>
  <style>
    body {{ font-family: "SimSun", serif; margin: 0; background: #f2f2f2; color: #111; }}
    .page {{ width: 210mm; min-height: 297mm; margin: 0 auto; box-sizing: border-box; padding: 22mm 18mm 22mm; background: white; }}
    h1,h2,h3,h4 {{ color: #111; }}
    h1 {{ text-align: center; font-size: 28px; margin: 0 0 18px; letter-spacing: 1px; }}
    .subtitle {{ display:none; }}
    section {{ background: white; border: none; padding: 0; margin: 0 0 18px; box-shadow: none; border-radius: 0; }}
    table {{ width:100%; border-collapse: collapse; margin: 12px 0; }}
    th,td {{ border: 1px solid #666; padding: 8px 10px; vertical-align: top; font-size: 14px; line-height: 1.7; }}
    th {{ background: #f0f0f0; }}
    p, li {{ font-size: 15px; line-height: 1.9; margin: 0 0 8px; }}
    ol, ul {{ margin: 8px 0 8px 24px; }}
    pre {{ background: #fafafa; color: #111; padding: 12px; border: 1px solid #999; overflow:auto; font-family: "SimSun", serif; font-size: 12px; line-height: 1.55; white-space: pre-wrap; word-break: break-all; }}
    .note, .tip, .warning {{ border: 1px solid #888; padding: 10px 12px; margin: 12px 0; background: #fafafa; }}
    .meta {{ color:#333; font-size: 13px; margin: 6px 0 10px; }}
    .figure img {{ width:100%; border:1px solid #999; }}
    .figure--module {{ margin-top: 14px; padding: 10px; border: 1px solid #999; background: #fff; }}
    .figcaption {{ text-align:center; font-size:13px; margin-top:6px; color:#111; }}
    .module + .module {{ margin-top: 20px; }}
    footer {{ text-align:center; color:#111; font-size:13px; margin-top:28px; }}
  </style>
</head>
<body>
  <div class="page">
    <h1>{esc(system_name)} 详细使用说明书</h1>

    <section>
      <h2>1. 系统简介</h2>
      <p><strong>系统定位：</strong>{esc(manifest.get("purpose", ""))}</p>
      <p><strong>面向用户群体：</strong>{esc(join_terms(manifest.get("user_groups", [])))}</p>
      <p><strong>总体架构：</strong>{esc(manifest.get("architecture", "前后端分离 B/S 架构，前端负责界面与交互，后端负责业务逻辑、接口与数据输出。"))}</p>
    <p><strong>技术方向：</strong>{esc(feature_categories)}</p>
    </section>

    <section>
     <h2>2. 系统运行环境要求</h2>
     {html_table(["项目", "内容"], env_rows)}
     <h3>2.1 软件技术特点分类</h3>
     <p><strong>{esc(feature_categories)}</strong></p>
     <h3>2.2 软件技术特点说明</h3>
     <p>{esc(feature_description)}</p>
    </section>

    <section>
      <h2>3. 系统安装与启动</h2>
      <ol>
        {''.join(f'<li>{esc(step)}</li>' for step in install_steps)}
      </ol>
      {f'<div class="warning">{esc(install_note)}</div>' if install_note else ''}
    </section>

    <section>
      <h2>4. 系统操作指南与系统的详细操作步骤</h2>
      {''.join(operation_modules)}
    </section>

    <section>
      <h2>5. 功能模块详解与对应主程序段</h2>
      {''.join(detail_modules)}
    </section>

    <section>
      <h2>6. 数据管理与导出</h2>
      <p>{esc(export_intro)}</p>
      {html_table(["导出文件", "字段说明", "使用方式"], export_rows)}
    </section>

    <section>
      <h2>7. 常见问题与故障排除</h2>
      {html_table(["问题现象", "可能原因", "排查步骤", "解决方案"], faq_rows)}
    </section>

    <section>
      <h2>8. 测试与维护</h2>
      <p><strong>内置测试：</strong>{esc(manifest.get("tests", {}).get("built_in", ""))}</p>
      <h3>8.1 测试用例</h3>
      {html_table(["编号", "测试内容"], test_case_rows)}
      <h3>8.2 系统维护指南</h3>
      <p>{esc(manifest.get("maintenance", {}).get("guide", ""))}</p>
      {maintenance_html}
    </section>

    <footer></footer>
  </div>
</body>
</html>"""
    output_path.write_text(html_doc, encoding="utf-8")


def fill_cell_text(cell, text: str, bold: bool = False, size: float = 12) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_font(run, size, bold=bold)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    add_border(cell)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, 18, bold=True)


def add_heading(doc: Document, level: int, text: str) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(8)
    fmt.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, 14 if level == 1 else 13, bold=True)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_body_paragraph(p)
    run = p.add_run(text)
    set_font(run, 12)


def add_code(doc: Document, label: str, code: str) -> None:
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    run = p.add_run(label + "\n" + code)
    set_font(run, 10.5)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    set_font(run, 10.5)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    escaped_docx = str(docx_path).replace("'", "''")
    escaped_pdf = str(pdf_path).replace("'", "''")
    script = (
        "$ErrorActionPreference='Stop';"
        "$word=New-Object -ComObject Word.Application;"
        "$word.Visible=$false;"
        f"$doc=$word.Documents.Open('{escaped_docx}');"
        f"$doc.SaveAs([ref]'{escaped_pdf}', [ref]17);"
        "$doc.Close();"
        "$word.Quit();"
    )
    try:
        subprocess.run(
            ["powershell", "-NoProfile", "-Command", script],
            check=True,
            capture_output=True,
            text=True,
        )
        return True
    except subprocess.CalledProcessError:
        return False


def build_docx(project_root: Path, manifest: dict, output_path: Path) -> None:
    system_name = manifest["system_name"]
    version = manifest.get("version", "V1.0")
    env = manifest["environments"]
    doc = Document()
    install_steps, install_note = build_install_steps(project_root, manifest)
    export_intro = build_export_intro(manifest)
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(12)

    add_title(doc, f"{system_name} 详细使用说明书")

    add_heading(doc, 1, "1. 系统简介")
    add_body(doc, f"系统定位：{manifest.get('purpose', '')}")
    add_body(doc, f"面向用户群体：{join_terms(manifest.get('user_groups', []))}")
    add_body(doc, f"总体架构：{manifest.get('architecture', '采用桌面业务架构，界面层、规则层和导入导出层协同运行。')}")
    feature_categories = infer_feature_categories(manifest)
    feature_description = build_feature_description(manifest)
    feature_table_text = build_feature_table_text(manifest)
    add_body(doc, f"技术方向：{feature_categories}")

    add_heading(doc, 1, "2. 系统运行环境要求")
    env_rows = [
        ("开发该软件的硬件环境", env.get("dev_hardware", "")),
        ("运行的硬件环境", env.get("run_hardware", "")),
        ("开发该软件的操作系统", env.get("dev_os", "")),
        ("软件开发环境/开发工具", env.get("dev_tools", "")),
        ("该软件的运行平台/操作系统", env.get("run_platform", "")),
        ("软件运行支撑环境/支持软件", env.get("support_env", "")),
        ("编程语言", join_terms(env.get("languages", []))),
        ("开发目的", manifest.get("purpose", "")),
        ("面向领域/行业", manifest.get("industry", "")),
        ("软件的主要功能", join_sentences(manifest.get("main_functions", []))),
        ("软件的技术特点", feature_table_text),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    fill_cell_text(table.rows[0].cells[0], "项目", bold=True)
    fill_cell_text(table.rows[0].cells[1], "内容", bold=True)
    for left, right in env_rows:
        row = table.add_row().cells
        fill_cell_text(row[0], left)
        fill_cell_text(row[1], right)

    add_heading(doc, 2, "2.1 软件技术特点分类")
    add_body(doc, feature_categories)
    add_heading(doc, 2, "2.2 软件技术特点说明")
    add_body(doc, feature_description)

    add_heading(doc, 1, "3. 系统安装与启动")
    for step in install_steps:
        add_body(doc, step)
    if install_note:
        add_body(doc, f"注意：{install_note}")

    add_heading(doc, 1, "4. 系统操作指南与系统的详细操作步骤")
    modules = manifest.get("modules", [])
    for index, module in enumerate(modules, start=1):
        add_heading(doc, 2, f"4.{index} {module['title']}")
        add_body(doc, f"菜单：{module.get('menu', '')}；页面路径：{module.get('route', '')}")
        for step_index, step in enumerate(module.get("steps", []), start=1):
            step_text = step if isinstance(step, str) else step.get("text", "")
            add_body(doc, f"第 {step_index} 步：{step_text}")
        for shot_index, shot in enumerate(collect_module_screenshots(module), start=1):
            shot_path = project_root / shot["file"]
            if shot_path.exists():
                doc.add_picture(io.BytesIO(build_manual_image_bytes(shot_path)), width=Cm(17.0))
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                label = f"图4-{index}-{shot_index}" if shot_index > 1 else f"图4-{index}"
                run = caption.add_run(f"{label} {shot.get('title', module['title'])}")
                set_font(run, 10.5)

    add_heading(doc, 1, "5. 功能模块详解与对应主程序段")
    for index, module in enumerate(manifest.get("modules", []), start=1):
        add_heading(doc, 2, f"5.{index} {module['title']}")
        add_body(doc, f"功能概述：{module.get('summary', '')}")
        add_body(doc, f"算法原理：{module.get('algorithm', '')}")
        add_body(doc, f"输入输出说明：{module.get('input_output', '')}")
        add_body(doc, f"数据流：{module.get('data_flow', '')}")
        add_body(doc, f"特殊处理：{module.get('special_handling', '')}")

        for ref_index, ref in enumerate(module.get("code_refs", []), start=1):
            label, snippet = code_slice(project_root, ref)
            add_heading(doc, 2, f"5.{index}.{ref_index} {ref.get('title', ref.get('function_name', '核心代码'))}")
            add_body(doc, f"源代码位置：{label}")
            add_body(doc, f"函数名：{ref.get('function_name', '')}")
            add_body(doc, f"说明：{ref.get('explanation', '')}")
            add_code(doc, "核心代码段：", snippet)
            add_body(doc, f"输入：{ref.get('input', '')}")
            add_body(doc, f"输出：{ref.get('output', '')}")
            if ref.get("variables"):
                var_table = doc.add_table(rows=1, cols=3)
                fill_cell_text(var_table.rows[0].cells[0], "变量名", bold=True)
                fill_cell_text(var_table.rows[0].cells[1], "含义", bold=True)
                fill_cell_text(var_table.rows[0].cells[2], "取值范围", bold=True)
                for item in ref["variables"]:
                    row = var_table.add_row().cells
                    fill_cell_text(row[0], item.get("name", ""))
                    fill_cell_text(row[1], item.get("meaning", ""))
                    fill_cell_text(row[2], item.get("range", ""))
            for item in ref.get("line_explanations", []):
                add_body(doc, f"{item.get('lines', '')}：{item.get('explanation', '')}")

    add_heading(doc, 1, "6. 数据管理与导出")
    add_body(doc, export_intro)
    export_table = doc.add_table(rows=1, cols=3)
    fill_cell_text(export_table.rows[0].cells[0], "导出文件", bold=True)
    fill_cell_text(export_table.rows[0].cells[1], "字段说明", bold=True)
    fill_cell_text(export_table.rows[0].cells[2], "使用方式", bold=True)
    for item in manifest.get("exports", []):
        row = export_table.add_row().cells
        fill_cell_text(row[0], item.get("name", ""))
        fill_cell_text(row[1], join_field_meanings(item.get("fields", [])))
        fill_cell_text(row[2], item.get("usage", ""))

    add_heading(doc, 1, "7. 常见问题与故障排除")
    faq_table = doc.add_table(rows=1, cols=4)
    for idx, title in enumerate(["问题现象", "可能原因", "排查步骤", "解决方案"]):
        fill_cell_text(faq_table.rows[0].cells[idx], title, bold=True)
    faq_items = manifest.get("faq", [])
    for item in faq_items:
        row = faq_table.add_row().cells
        fill_cell_text(row[0], item.get("problem", ""))
        fill_cell_text(row[1], item.get("reason", ""))
        fill_cell_text(row[2], item.get("steps", ""))
        fill_cell_text(row[3], item.get("solution", ""))

    add_heading(doc, 1, "8. 测试与维护")
    add_body(doc, f"内置测试：{manifest.get('tests', {}).get('built_in', '')}")

    add_heading(doc, 2, "8.1 测试用例")
    test_case_table = doc.add_table(rows=1, cols=2)
    fill_cell_text(test_case_table.rows[0].cells[0], "编号", bold=True)
    fill_cell_text(test_case_table.rows[0].cells[1], "测试内容", bold=True)
    test_cases = manifest.get("tests", {}).get("cases", [])
    for index, case in enumerate(test_cases, start=1):
        row = test_case_table.add_row().cells
        fill_cell_text(row[0], f"TC-{index:02d}")
        fill_cell_text(row[1], case)

    add_heading(doc, 2, "8.2 系统维护指南")
    add_body(doc, manifest.get('maintenance', {}).get('guide', ''))
    maintenance_items = manifest.get("maintenance", {}).get("items", [])
    for item in maintenance_items:
        add_body(doc, f"维护事项：{item.get('title', '')}")
        add_body(doc, f"对应位置：{item.get('target', '')}")
        add_body(doc, f"处理方法：{item.get('method', '')}")
        add_body(doc, f"完成后验证：{item.get('verify', '')}")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_page = doc.sections[0].footer.add_paragraph()
    footer_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_page)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate HTML and DOCX software copyright manual from manifest.")
    parser.add_argument("--project", required=True, help="Project root path")
    parser.add_argument("--manifest", required=True, help="Manifest JSON path")
    parser.add_argument("--output-root", help="Output directory, defaults to <project>/软件著作权申请资料/正式资料")
    parser.add_argument("--skip-pdf", action="store_true", help="Do not emit PDF for the manual")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    project_root = Path(args.project).resolve()
    manifest_path = Path(args.manifest).resolve()
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    output_root = Path(args.output_root).resolve() if args.output_root else project_root / "软件著作权申请资料" / "正式资料"
    system_name = manifest["system_name"]

    html_out = output_root / f"{system_name}_软件著作权鉴定材料.html"
    docx_out = output_root / f"{system_name}_软件著作权鉴定材料.docx"
    pdf_out = output_root / f"{system_name}_软件著作权鉴定材料.pdf"
    build_html(project_root, manifest, html_out)
    build_docx(project_root, manifest, docx_out)
    pdf_created = False
    if not args.skip_pdf:
        pdf_created = convert_docx_to_pdf(docx_out, pdf_out)
    print(f"HTML={html_out}")
    print(f"DOCX={docx_out}")
    print(f"PDF={pdf_out if pdf_created else 'NOT_CREATED'}")


if __name__ == "__main__":
    main()
