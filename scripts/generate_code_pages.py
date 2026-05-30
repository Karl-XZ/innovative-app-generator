from __future__ import annotations

import argparse
import json
import math
import re
import subprocess
import tempfile
import unicodedata
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pypdf import PdfReader, PdfWriter


PAGE_LINE_TARGET = 50
FRONT_BACK_PAGES = 30
CODE_FONT_SIZE = 10.5  # 五号
INITIAL_LINE_SPACING = 14.0
MAX_WRAP_UNITS = 88

INCLUDE_EXTENSIONS = {
    ".java",
    ".kt",
    ".py",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".css",
    ".scss",
    ".less",
    ".html",
    ".vue",
    ".xml",
    ".properties",
    ".sql",
    ".json",
    ".yaml",
    ".yml",
}

EXCLUDED_DIRS = {
    ".git",
    ".idea",
    ".vscode",
    ".next",
    ".nuxt",
    ".cache",
    ".turbo",
    ".parcel-cache",
    ".pytest_cache",
    "__pycache__",
    "node_modules",
    "dist",
    "build",
    "out",
    "target",
    "coverage",
    "vendor",
    "tmp",
    "temp",
    "docs",
    "screenshots",
    "软件著作权申请资料",
}

EXCLUDED_PATH_PREFIXES = (
    "scripts/",
    "docs/",
    "screenshots/",
    "软件著作权申请资料/",
)

EXCLUDED_FILENAMES = {
    "softcopyright-manifest.json",
    "package-lock.json",
    "package.json",
    "tsconfig.json",
    "tsconfig.app.json",
    "tsconfig.node.json",
    "eslint.config.js",
    "vite.config.ts",
}

PRIORITY_PARTS = [
    ("java-server", "src"),
    ("src",),
    ("app",),
    ("server",),
    ("backend",),
    ("frontend",),
    ("java-server", "resources"),
    ("resources",),
    ("public",),
]


@dataclass(frozen=True)
class SourceFile:
    path: Path
    relative_path: str
    category: str


def set_run_font(run, size: float, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = False


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def count_lines(path: Path) -> int:
    return len(read_text(path).splitlines())


def infer_category(relative_path: str) -> str:
    lower = relative_path.lower()
    if lower.endswith((".java", ".kt")):
        return "后端源码"
    if lower.endswith((".ts", ".tsx", ".js", ".jsx", ".vue")):
        return "前端脚本"
    if lower.endswith((".css", ".scss", ".less")):
        return "界面样式"
    if lower.endswith(".html"):
        return "页面模板"
    if lower.endswith((".xml", ".properties", ".yaml", ".yml", ".json", ".sql")):
        return "配置与数据脚本"
    return "项目源码"


def is_excluded(path: Path, project_root: Path) -> bool:
    relative = path.relative_to(project_root)
    relative_posix = relative.as_posix()
    if path.name in EXCLUDED_FILENAMES:
        return True
    if any(relative_posix.startswith(prefix) for prefix in EXCLUDED_PATH_PREFIXES):
        return True
    return any(part in EXCLUDED_DIRS for part in relative.parts)


def priority_rank(relative_parts: tuple[str, ...]) -> int:
    lowered = tuple(part.lower() for part in relative_parts)
    for index, prefix in enumerate(PRIORITY_PARTS):
        if lowered[: len(prefix)] == prefix:
            return index
    return len(PRIORITY_PARTS)


def collect_source_files(project_root: Path) -> list[SourceFile]:
    files: list[SourceFile] = []
    for path in project_root.rglob("*"):
        if not path.is_file():
            continue
        if path.suffix.lower() not in INCLUDE_EXTENSIONS:
            continue
        if is_excluded(path, project_root):
            continue
        relative = path.relative_to(project_root).as_posix()
        files.append(SourceFile(path=path, relative_path=relative, category=infer_category(relative)))
    files.sort(key=lambda item: (priority_rank(tuple(item.path.relative_to(project_root).parts)), item.relative_path.lower()))
    return files


def display_units(text: str) -> int:
    total = 0
    for char in text:
        total += 2 if unicodedata.east_asian_width(char) in {"W", "F"} else 1
    return total


def wrap_content_line(line: str, max_units: int = MAX_WRAP_UNITS) -> list[str]:
    if not line:
        return [""]
    wrapped: list[str] = []
    remaining = line
    continuation_prefix = "    "
    while display_units(remaining) > max_units:
        units = 0
        split_at = 0
        for index, char in enumerate(remaining, start=1):
            units += 2 if unicodedata.east_asian_width(char) in {"W", "F"} else 1
            if units > max_units:
                break
            split_at = index
        if split_at <= 0:
            split_at = min(len(remaining), max_units)
        wrapped.append(remaining[:split_at])
        remaining = continuation_prefix + remaining[split_at:]
    wrapped.append(remaining)
    return wrapped


def build_content_lines(system_name: str, files: list[SourceFile]) -> list[str]:
    lines: list[str] = []
    lines.extend(
        [
            f"【软件名称】{system_name}",
            "【代码说明】以下内容为完整项目源程序连续清单",
            "",
        ]
    )
    for source in files:
        lines.extend(
            wrap_content_line(f"【文件】{source.relative_path}")
            + wrap_content_line(f"【类型】{source.category}")
        )
        for raw_line in read_text(source.path).splitlines():
            lines.extend(wrap_content_line(raw_line))
        lines.append("")
    return lines


def expected_page_count(content_lines: list[str]) -> int:
    return math.ceil(len(content_lines) / PAGE_LINE_TARGET)


def build_docx(system_name: str, content_lines: list[str], output_path: Path, line_spacing_pt: float) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(CODE_FONT_SIZE)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run(f"{system_name}代码")
    set_run_font(header_run, 10.5, bold=False)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer)

    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(0)
    fmt.line_spacing = Pt(line_spacing_pt)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    run = paragraph.add_run("\n".join(content_lines))
    set_run_font(run, CODE_FONT_SIZE)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, 10.5)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    escaped_docx = str(docx_path).replace("'", "''")
    escaped_pdf = str(pdf_path).replace("'", "''")
    script = (
        "$ErrorActionPreference='Stop';"
        "$word=New-Object -ComObject Word.Application;"
        "$word.Visible=$false;"
        f"$doc=$word.Documents.Open('{escaped_docx}');"
        f"$doc.SaveAs([ref]'{escaped_pdf}', [ref]17);"
        "$doc.Close();"
        "$word.Quit();"
    )
    try:
        subprocess.run(["powershell", "-NoProfile", "-Command", script], check=True, capture_output=True, text=True)
        return True
    except subprocess.CalledProcessError:
        return False


def count_pdf_pages(pdf_path: Path) -> int:
    return len(PdfReader(str(pdf_path)).pages)


def pick_front_back_pdf_pages(full_pdf: Path, final_pdf: Path) -> int:
    reader = PdfReader(str(full_pdf))
    total = len(reader.pages)
    writer = PdfWriter()
    if total <= FRONT_BACK_PAGES * 2:
        indexes = list(range(total))
    else:
        indexes = list(range(FRONT_BACK_PAGES)) + list(range(total - FRONT_BACK_PAGES, total))
    for index in indexes:
        writer.add_page(reader.pages[index])
    with final_pdf.open("wb") as fh:
        writer.write(fh)
    return len(indexes)


def calibrate_line_spacing(system_name: str, content_lines: list[str], full_docx: Path) -> tuple[float, Path, int]:
    expected_pages = expected_page_count(content_lines)
    temp_dir = Path(tempfile.mkdtemp(prefix="code-pages-"))
    best_pdf = temp_dir / "full.pdf"
    best_actual = -1
    chosen_spacing = INITIAL_LINE_SPACING
    low = 12.0
    high = 16.0
    best_diff = None

    for _ in range(8):
        spacing = round((low + high) / 2, 2)
        build_docx(system_name, content_lines, full_docx, spacing)
        if not convert_docx_to_pdf(full_docx, best_pdf):
            raise SystemExit("代码 Word 转 PDF 失败，无法校验页数。")
        actual_pages = count_pdf_pages(best_pdf)
        if actual_pages == expected_pages:
            return spacing, best_pdf, actual_pages
        diff = abs(actual_pages - expected_pages)
        if best_diff is None or diff < best_diff:
            best_diff = diff
            best_actual = actual_pages
            chosen_spacing = spacing
        if actual_pages < expected_pages:
            low = spacing
        else:
            high = spacing
    if best_actual != expected_pages:
        raise SystemExit(f"代码量与页数校验失败：内容应为 {expected_pages} 页，实际导出为 {best_actual} 页。")
    return chosen_spacing, best_pdf, best_actual


def clip_text(text: str, limit: int) -> str:
    text = " ".join((text or "").split())
    return text[:limit]


def infer_languages(source_files: list[SourceFile]) -> list[str]:
    mapping = {
        ".java": "Java",
        ".kt": "Kotlin",
        ".ts": "TypeScript",
        ".tsx": "TypeScript",
        ".js": "JavaScript",
        ".jsx": "JavaScript",
        ".html": "HTML",
        ".css": "CSS",
        ".scss": "CSS",
        ".less": "CSS",
        ".xml": "XML",
        ".json": "JSON",
    }
    ordered_suffixes = [".java", ".kt", ".ts", ".tsx", ".js", ".jsx", ".html", ".css", ".scss", ".less", ".xml", ".json"]
    found = {item.path.suffix.lower() for item in source_files}
    result: list[str] = []
    for suffix in ordered_suffixes:
        if suffix in found:
            label = mapping[suffix]
            if label not in result:
                result.append(label)
    return result


def read_optional_text(path: Path) -> str:
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="replace")


def infer_environment(project_root: Path, manifest: dict, source_files: list[SourceFile]) -> dict:
    env = dict(manifest.get("environments", {}))
    languages = env.get("languages") or infer_languages(source_files)
    has_java = "Java" in languages or "Kotlin" in languages
    has_node = (project_root / "package.json").exists()
    env.setdefault("dev_hardware", "Intel i5及以上CPU，16GB内存，50GB硬盘")
    env.setdefault("run_hardware", "双核CPU，4GB内存，2GB可用磁盘")
    env.setdefault("dev_os", "Windows 10/11 64位")
    if not env.get("dev_tools"):
        tools = []
        if has_java:
            tools.append("JDK 21")
        if has_node:
            tools.extend(["Node.js", "Vite"])
        tools.append("Visual Studio Code")
        env["dev_tools"] = "、".join(tools)
    env.setdefault("run_platform", "Windows、macOS、Linux")
    if not env.get("support_env"):
        supports = []
        if has_java:
            supports.append("JDK 21")
        if has_node:
            supports.append("现代浏览器")
        env["support_env"] = "、".join(supports) if supports else "现代浏览器"
    env["languages"] = languages
    return env


def infer_purpose_and_industry(system_name: str, manifest: dict, project_root: Path) -> tuple[str, str]:
    purpose = manifest.get("purpose", "")
    industry = manifest.get("industry", "")
    readme = read_optional_text(project_root / "README.md")
    if not purpose:
        if "医院" in system_name or "医疗" in readme:
            purpose = "互联网医院诊疗与运营协同管理"
        else:
            purpose = system_name
    if not industry:
        if "医院" in system_name or "医疗" in readme:
            industry = "互联网医疗、智慧医院管理"
        else:
            industry = "行业业务管理软件"
    return clip_text(purpose, 50), clip_text(industry, 50)


def infer_feature_categories(manifest: dict, system_name: str, industry: str) -> str:
    categories = manifest.get("technical_feature_categories")
    if categories:
        return "、".join(categories[:2])
    if "医疗" in industry or "医院" in system_name:
        return "人工智能软件、医疗软件"
    if "教育" in industry:
        return "人工智能软件、教育软件"
    return "人工智能软件、应用软件"


def infer_main_functions(project_root: Path, manifest: dict) -> list[str]:
    if manifest.get("main_functions"):
        return list(manifest["main_functions"])
    readme = read_optional_text(project_root / "README.md")
    features: list[str] = []
    capture = False
    for raw_line in readme.splitlines():
        line = raw_line.strip()
        if line.startswith("## ") and "主要特性" in line:
            capture = True
            continue
        if capture and line.startswith("## "):
            break
        if capture and line.startswith("- "):
            features.append(line[2:].strip())
    return features


def build_main_functions_text(project_root: Path, manifest: dict) -> str:
    items = infer_main_functions(project_root, manifest)
    if not items:
        items = [module.get("summary", module.get("title", "")) for module in manifest.get("modules", []) if module.get("summary") or module.get("title")]
    if not items:
        items = ["系统提供完整业务流程管理、数据处理、结果展示和导出归档能力。"]
    parts = [f"（{idx}）{item}" for idx, item in enumerate(items, start=1)]
    extra_modules = manifest.get("modules", [])
    for module in extra_modules:
        text = module.get("summary") or module.get("algorithm") or ""
        if text:
            candidate = f"（{len(parts)+1}）{text}"
            if len("".join(parts) + candidate) <= 1200:
                parts.append(candidate)
        if len("".join(parts)) >= 600:
            break
    combined = "".join(parts)
    if len(combined) < 600:
        supplement = "系统支持完整中文界面、模块联动操作、结果导出、异常提示、数据校验与参数维护，可满足实际业务运行、成果归档与后续扩展需要。"
        while len(combined) < 600:
            candidate = f"（{len(parts)+1}）{supplement}"
            if len(combined) + len(candidate) > 1200:
                break
            parts.append(candidate)
            combined = "".join(parts)
    combined = "".join(parts)
    return combined[:1200]


def build_tech_features_text(manifest: dict) -> str:
    raw_items = manifest.get("technical_features", [])
    cleaned_items: list[str] = []
    for item in raw_items:
        text = str(item).strip()
        if not text:
            continue
        text = re.sub(r"[。；;、，,\s]+$", "", text)
        cleaned_items.append(text)
    if not cleaned_items:
        cleaned_items = ["采用模块化架构、规则引擎与结构化数据导出机制，支持业务联动、结果展示与资料生成"]

    features = ""
    for item in cleaned_items:
        candidate = item if not features else f"{features}；{item}"
        if len(candidate) <= 100:
            features = candidate
        else:
            break

    if not features:
        features = clip_text(cleaned_items[0], 100).rstrip("；、，,。 ")
    return features


def build_application_info(system_name: str, version: str, manifest: dict, total_lines: int, project_root: Path, source_files: list[SourceFile]) -> str:
    env = infer_environment(project_root, manifest, source_files)
    purpose, industry = infer_purpose_and_industry(system_name, manifest, project_root)
    fields = [
        ("软件全称", system_name),
        ("版本号", version),
        ("软件分类", manifest.get("software_category", "应用软件")),
        ("开发的硬件环境", clip_text(env.get("dev_hardware", ""), 50)),
        ("运行的硬件环境", clip_text(env.get("run_hardware", ""), 50)),
        ("开发该软件的操作系统", clip_text(env.get("dev_os", ""), 50)),
        ("软件开发环境/开发工具", clip_text(env.get("dev_tools", ""), 50)),
        ("该软件的运行平台/操作系统", clip_text(env.get("run_platform", ""), 50)),
        ("软件运行支撑环境/支持软件", clip_text(env.get("support_env", ""), 50)),
        ("编程语言", clip_text("、".join(env.get("languages", [])), 50)),
        ("源程序量", f"{total_lines}行"),
        ("开发目的", purpose),
        ("面向领域/行业", industry),
        ("软件的主要功能", build_main_functions_text(project_root, manifest)),
        ("软件的技术特点", f"{infer_feature_categories(manifest, system_name, industry)}；{build_tech_features_text(manifest)}"),
    ]
    return "\n".join(f"【{label}】{value}" for label, value in fields)


def load_manifest(path: Path | None) -> dict:
    if path and path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate full code DOCX and front/back 30-page PDF for software copyright use.")
    parser.add_argument("--project", required=True, help="Project root path")
    parser.add_argument("--manifest", help="Manifest JSON path")
    parser.add_argument("--system-name", help="Software name shown in the documents")
    parser.add_argument("--version", help="Software version")
    parser.add_argument("--output-root", help="Output directory, defaults to <project>/软件著作权申请资料/正式资料")
    parser.add_argument("--skip-pdf", action="store_true", help="Do not emit PDF for the code document")
    parser.add_argument("--skip-app-info", action="store_true", help="Do not emit XXX.txt")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    project_root = Path(args.project).resolve()
    manifest = load_manifest(Path(args.manifest).resolve() if args.manifest else (project_root / "softcopyright-manifest.json"))

    system_name = args.system_name or manifest.get("system_name") or project_root.name
    version = args.version or manifest.get("version") or "V1.0"
    output_root = Path(args.output_root).resolve() if args.output_root else project_root / "软件著作权申请资料" / "正式资料"

    source_files = collect_source_files(project_root)
    if not source_files:
        raise SystemExit("未找到符合规则的源码文件。")

    content_lines = build_content_lines(system_name, source_files)
    total_source_lines = sum(count_lines(item.path) for item in source_files)

    temp_full_docx = output_root / "__code_temp__.docx"
    info_txt = output_root / f"{system_name}.txt"

    spacing, temp_full_pdf, actual_full_pages = calibrate_line_spacing(system_name, content_lines, temp_full_docx)
    expected_pages = expected_page_count(content_lines)
    if actual_full_pages != expected_pages:
        raise SystemExit(f"Code page validation failed: expected {expected_pages} pages, got {actual_full_pages}.")

    if actual_full_pages > FRONT_BACK_PAGES * 2:
        full_docx = output_root / f"{system_name}代码（完整版）.docx"
        final_pdf = output_root / f"{system_name}代码（60页）.pdf"
    else:
        full_docx = output_root / f"{system_name}代码.docx"
        final_pdf = output_root / f"{system_name}代码.pdf"

    if full_docx.exists():
        full_docx.unlink()
    temp_full_docx.replace(full_docx)

    final_pdf_pages = 0
    if not args.skip_pdf:
        if final_pdf.exists():
            final_pdf.unlink()
        final_pdf_pages = pick_front_back_pdf_pages(temp_full_pdf, final_pdf)
        if actual_full_pages > FRONT_BACK_PAGES * 2 and final_pdf_pages != FRONT_BACK_PAGES * 2:
            raise SystemExit(f"Code PDF page validation failed: expected 60 pages, got {final_pdf_pages}.")

    if not args.skip_app_info:
        info_txt.write_text(build_application_info(system_name, version, manifest, total_source_lines, project_root, source_files), encoding="utf-8")

    print(f"FULL_DOCX={full_docx}")
    print(f"FINAL_PDF={final_pdf if final_pdf_pages else 'NOT_CREATED'}")
    print(f"INFO_TXT={info_txt if not args.skip_app_info else 'SKIPPED'}")
    print(f"TOTAL_SOURCE_LINES={total_source_lines}")
    print(f"CONTENT_LINES={len(content_lines)}")
    print(f"EXPECTED_PAGES={expected_pages}")
    print(f"ACTUAL_PAGES={actual_full_pages}")
    print(f"LINE_SPACING_PT={spacing}")
    print(f"TOTAL_SOURCE_LINES={total_source_lines}")
    print(f"CONTENT_LINES={len(content_lines)}")
    print(f"EXPECTED_PAGES={expected_pages}")
    print(f"ACTUAL_PAGES={actual_full_pages}")
    print(f"LINE_SPACING_PT={spacing}")


if __name__ == "__main__":
    main()
